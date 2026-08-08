#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Konversi docx Penjelasan RPJMD-Pemdi -> PDF rapi."""
import docx, os
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.units import mm

src = 'public/bukti-dukung/TataKelola_I1_24_Penjelasan-RPJMD-Pemdi_2025.docx'
dst = 'public/bukti-dukung/final-src/I1_Penjelasan_RPJMD.pdf'

doc = docx.Document(src)
paras = []
for p in doc.paragraphs:
    t = p.text.strip()
    if t:
        style = p.style.name if p.style else ''
        paras.append((style, t))
for tbl in doc.tables:
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        paras.append(('Table', ' | '.join(cells)))

styles = getSampleStyleSheet()
h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=13, spaceAfter=8)
body = ParagraphStyle('Body', parent=styles['BodyText'], fontSize=10.5, leading=15, spaceAfter=6)

story = [Paragraph('Penjelasan Substansi RAN Pemdi pada Dokumen Perencanaan (RPJMK Aceh Tengah 2025-2029)', h1)]
for style, t in paras:
    if 'Heading' in style:
        story.append(Paragraph(t, h1))
    elif style == 'Table':
        story.append(Paragraph(t.replace('|', ' • '), body))
    else:
        t = t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(t, body))

pdf = SimpleDocTemplate(dst, pagesize=A4, leftMargin=18*mm, rightMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm)
pdf.build(story)
print(f"OK {dst} {os.path.getsize(dst)//1024}K")
